import streamlit as st
import pandas as pd
import google.generativeai as genai
import os
import io
from datetime import datetime
from docx import Document

# ----------------- CONFIG -----------------
st.set_page_config(page_title="🎓 Smart AI Classroom", layout="wide")

# Hide Streamlit deploy/share buttons
st.markdown("""
<style>
#MainMenu, footer {visibility: hidden;}
.stDeployButton, .stAppDeployButton {display: none !important;}
</style>
""", unsafe_allow_html=True)

# ----------------- GOOGLE GEMINI -----------------
# 🔑 Replace with your own Gemini API key from Google AI Studio
genai.configure(api_key="AIzaSyCaoHHc8yG8Sg9_A96D4B_0mIuws1pAreI")  # <---- PUT YOUR KEY HERE
model = genai.GenerativeModel("gemini-1.5-flash")

# ----------------- DIRECTORIES -----------------
USER_FILE = "users.csv"
HISTORY_DIR = "student_history"
os.makedirs(HISTORY_DIR, exist_ok=True)

# ----------------- FUNCTIONS -----------------
def get_gemini_reply(prompt: str) -> str:
    """Generate friendly AI tutor answer"""
    try:
        if not prompt.strip():
            prompt = "Please clarify your question."
        system_prompt = (
            "You are a friendly AI tutor helping students during online class. "
            "Answer clearly, conversationally, and correctly. "
            "If the question is unclear, restate it before answering."
        )
        response = model.generate_content(f"{system_prompt}\n\nStudent asked: {prompt}\n\nAnswer:")
        return response.text.strip() if response and hasattr(response, "text") else "I couldn’t get that, please try again."
    except Exception as e:
        return f"⚠️ Error: {e}"

def save_chat(student, q, a):
    """Save student Q&A to CSV"""
    f = os.path.join(HISTORY_DIR, f"{student}_history.csv")
    new = pd.DataFrame([[datetime.now(), q, a]], columns=["Time", "Question", "Answer"])
    df = pd.read_csv(f) if os.path.exists(f) else pd.DataFrame(columns=["Time","Question","Answer"])
    pd.concat([df, new], ignore_index=True).to_csv(f, index=False)

def load_chat(student):
    """Load student chat history"""
    f = os.path.join(HISTORY_DIR, f"{student}_history.csv")
    return pd.read_csv(f) if os.path.exists(f) else pd.DataFrame(columns=["Time","Question","Answer"])

# ----------------- APP UI -----------------
st.title("🎓 Smart AI Classroom")

menu = st.sidebar.radio("Choose Role", ["Student", "Teacher"])

# ----------------- STUDENT PAGE -----------------
if menu == "Student":
    st.subheader("👩‍🎓 Student Class Page")

    student_name = st.text_input("Enter your name to join class:")

    if st.button("Join Class"):
        if student_name.strip():
            st.session_state.student = student_name
            st.success(f"✅ Welcome {student_name}! Attendance marked automatically.")
        else:
            st.warning("Please enter your name.")

    if "student" in st.session_state:
        # Show broadcast (if any)
        if st.session_state.get("broadcast", ""):
            st.info(f"📢 Announcement: {st.session_state.broadcast}")

        st.write("---")
        st.subheader("💬 Ask a Question")
        prompt = st.text_input("Your Question:")
        if st.button("Ask Gemini"):
            if prompt.strip():
                reply = get_gemini_reply(prompt)
                save_chat(st.session_state.student, prompt, reply)
                st.markdown(f"**🤖 Gemini:** {reply}")
            else:
                st.warning("Please type a question first.")

        st.write("---")
        st.subheader("📘 Your Previous Questions")
        chat_history = load_chat(st.session_state.student)
        st.dataframe(chat_history, use_container_width=True)

        if not chat_history.empty:
            # Create Word document in memory
            doc = Document()
            doc.add_heading(f"Chat History - {st.session_state.student}", 1)
            for _, row in chat_history.iterrows():
                doc.add_paragraph(f"🕒 {row['Time']}")
                doc.add_paragraph(f"Q: {row['Question']}")
                doc.add_paragraph(f"A: {row['Answer']}")
                doc.add_paragraph("---")

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Browser download button
            st.download_button(
                label="⬇️ Download My Chat (.docx)",
                data=buffer,
                file_name=f"{st.session_state.student}_Chat.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ----------------- TEACHER PAGE -----------------
elif menu == "Teacher":
    st.subheader("🧑‍🏫 Teacher Dashboard")

    teacher_name = st.text_input("Enter your name to access dashboard:")

    if st.button("Enter Dashboard"):
        if teacher_name.strip():
            st.session_state.teacher = teacher_name
            st.success(f"✅ Welcome, {teacher_name} (Teacher Dashboard)")
        else:
            st.warning("Please enter your name.")

    if "teacher" in st.session_state:
        st.write("---")
        st.subheader("📢 Broadcast Message to Students")
        msg = st.text_input("Enter announcement")
        if st.button("Send Broadcast"):
            st.session_state.broadcast = msg
            st.success("✅ Broadcast sent to all students!")

        st.write("---")
        st.subheader("📋 All Student Questions")

        all_data = []
        for file in os.listdir(HISTORY_DIR):
            student = file.replace("_history.csv", "")
            df = pd.read_csv(os.path.join(HISTORY_DIR, file))
            for _, r in df.iterrows():
                all_data.append([student, r.get("Time",""), r.get("Question",""), r.get("Answer","")])

        if all_data:
            df_all = pd.DataFrame(all_data, columns=["Student", "Time", "Question", "Answer"])
            st.dataframe(df_all, use_container_width=True)

            # Browser download for teacher summary
            doc = Document()
            doc.add_heading("All Student Q&A Summary", 0)
            for _, row in df_all.iterrows():
                doc.add_paragraph(f"👩‍🎓 {row['Student']} asked: {row['Question']}")
                doc.add_paragraph(f"🤖 Answer: {row['Answer']}")
                doc.add_paragraph("---")

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="⬇️ Download All Chats (.docx)",
                data=buffer,
                file_name="teacher_class_summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.write("---")
        if st.button("🆕 End Current Class"):
            st.session_state.clear()
            st.success("Class ended. Session cleared. ✅")
